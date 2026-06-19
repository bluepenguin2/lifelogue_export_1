# -*- coding: utf-8 -*-
"""
라이프로그 알림장 백업 - 백엔드 서버 (FastAPI)

실행:
    pip install fastapi uvicorn requests python-docx
    uvicorn main:app --host 0.0.0.0 --port 8000
"""

import io
import re
from calendar import monthrange
from datetime import date
from typing import Optional

import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

app = FastAPI()

# CORS 허용 (모바일 브라우저에서 직접 호출 가능하도록)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

API_BASE = "https://app.dailylogue.com/user-service/api/v1"
IMAGE_DOWNLOAD_BASE = "https://lifelogue.kr/user-service/api/v1/file/download/images"

CREATOR_TYPE_KR = {
    "TEACHER": "선생님",
    "MOTHER": "어머니",
    "FATHER": "아버지",
    "STUDENT": "원아",
}


# =========================================================
# 요청 모델
# =========================================================
class GenerateRequest(BaseModel):
    token: str          # "Bearer eyJ..." 형태
    start_year: int
    start_month: int


# =========================================================
# 유틸 함수
# =========================================================
def month_range(year, month):
    last_day = monthrange(year, month)[1]
    return f"{year:04d}-{month:02d}-01", f"{year:04d}-{month:02d}-{last_day:02d}"


def iter_year_months(start_year, start_month, end_year, end_month):
    y, m = start_year, start_month
    while (y, m) <= (end_year, end_month):
        yield y, m
        m = m + 1 if m < 12 else 1
        if m == 1:
            y += 1


def html_to_paragraphs(html_content):
    if not html_content:
        return []
    paragraphs = re.findall(r"<p[^>]*>(.*?)</p>", html_content, flags=re.DOTALL)
    if not paragraphs:
        paragraphs = [html_content]
    cleaned = []
    for p in paragraphs:
        p = re.sub(r"<br\s*/?>", "\n", p)
        p = re.sub(r"<[^>]+>", "", p)
        p = (p.replace("&nbsp;", " ").replace("&amp;", "&")
              .replace("&lt;", "<").replace("&gt;", ">")
              .replace("&quot;", '"').replace("&#39;", "'"))
        cleaned.append(p.strip())
    return cleaned


# =========================================================
# 라이프로그 API 호출
# =========================================================
def get_user_info(token):
    headers = {"Authorization": token}
    resp = requests.get(f"{API_BASE}/auth/me", headers=headers, timeout=10)
    resp.raise_for_status()
    data = resp.json()
    org = data["organizations"][0]
    child = org["children"][0]
    return {
        "organizationId": org["id"],
        "targetId": child["id"],
        "childName": child["name"],
        "orgName": org["name"],
    }


def fetch_notices_for_month(token, organization_id, target_id, year, month):
    start_date, end_date = month_range(year, month)
    headers = {"Authorization": token}
    all_items = []
    page = 0
    while True:
        params = {
            "organizationId": organization_id,
            "targetId": target_id,
            "startDate": start_date,
            "endDate": end_date,
            "page": page,
            "size": 12,
        }
        resp = requests.get(
            f"{API_BASE}/notice/list/child",
            headers=headers, params=params, timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        all_items.extend(data.get("content", []))
        if page + 1 >= data.get("totalPages", 1):
            break
        page += 1
    return all_items


def download_image(token, image_id):
    try:
        resp = requests.get(
            f"{IMAGE_DOWNLOAD_BASE}/{image_id}",
            headers={"Authorization": token},
            timeout=20,
        )
        resp.raise_for_status()
        return resp.content
    except Exception:
        return None


# =========================================================
# Word 문서 생성
# =========================================================
def build_document(notices_by_month, child_name, org_name, token):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{org_name}\n{child_name} 알림장 모음")
    run.font.size = Pt(20)
    run.font.bold = True
    doc.add_page_break()

    for year, month, notices in notices_by_month:
        if not notices:
            continue
        doc.add_heading(f"{year}년 {month}월", level=1)
        for notice in sorted(notices, key=lambda n: n.get("publishedAt", "")):
            add_notice_to_doc(doc, notice, token)
        doc.add_page_break()

    return doc


def add_notice_to_doc(doc, notice, token):
    title_text = notice.get("title") or "(제목 없음)"
    published_at = notice.get("publishedAt", "")
    date_str = published_at[:10] if published_at else ""
    creator = notice.get("creator") or {}
    creator_name = creator.get("name", "")
    creator_type_kr = CREATOR_TYPE_KR.get(creator.get("type", ""), "")

    p = doc.add_paragraph()
    run = p.add_run(f"[{date_str}] {title_text}")
    run.font.bold = True
    run.font.size = Pt(13)

    if creator_name:
        meta = doc.add_paragraph()
        r = meta.add_run(f"작성자: {creator_name} ({creator_type_kr})")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for para_text in html_to_paragraphs(notice.get("content", "")):
        doc.add_paragraph().add_run(para_text)

    for img in (notice.get("images") or []):
        image_bytes = download_image(token, img.get("id"))
        if image_bytes:
            try:
                doc.add_picture(io.BytesIO(image_bytes), width=Inches(4))
            except Exception:
                doc.add_paragraph().add_run("[이미지를 표시할 수 없습니다]").font.italic = True
        else:
            doc.add_paragraph().add_run("[이미지를 불러오지 못했습니다]").font.italic = True

    for comment in (notice.get("comments") or []):
        c_type_kr = CREATOR_TYPE_KR.get(comment.get("relationType", ""), "")
        c_p = doc.add_paragraph()
        c_p.paragraph_format.left_indent = Inches(0.3)
        r = c_p.add_run(f"└ ({c_type_kr} 댓글) {comment.get('content', '')}")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x66, 0x33)

    sep = doc.add_paragraph()
    sep.add_run("─" * 40).font.size = Pt(8)


# =========================================================
# API 엔드포인트
# =========================================================
@app.get("/", response_class=HTMLResponse)
async def root():
    """프론트엔드 HTML을 서빙한다"""
    import os
    html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "index.html")
    with open(html_path, encoding="utf-8") as f:
        return f.read()


@app.post("/api/user-info")
async def api_user_info(body: dict):
    """토큰으로 사용자/자녀 정보를 반환한다"""
    token = body.get("token", "")
    if not token:
        raise HTTPException(status_code=400, detail="token이 필요합니다")
    try:
        return get_user_info(token if token.startswith("Bearer ") else f"Bearer {token}")
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"사용자 정보 조회 실패: {e}")


@app.post("/api/generate")
async def api_generate(req: GenerateRequest):
    """알림장을 수집하여 Word 파일을 생성하고 반환한다"""
    token = req.token if req.token.startswith("Bearer ") else f"Bearer {req.token}"

    try:
        user_info = get_user_info(token)
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"인증 실패: {e}")

    today = date.today()
    notices_by_month = []
    for year, month in iter_year_months(
        req.start_year, req.start_month, today.year, today.month
    ):
        notices = fetch_notices_for_month(
            token, user_info["organizationId"], user_info["targetId"], year, month
        )
        notices_by_month.append((year, month, notices))

    doc = build_document(
        notices_by_month, user_info["childName"], user_info["orgName"], token
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (
        f"{user_info['childName']}_알림장_"
        f"{req.start_year}{req.start_month:02d}-"
        f"{today.year}{today.month:02d}.docx"
    )

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{requests.utils.quote(filename)}"},
    )
